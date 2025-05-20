import streamlit as st
import pandas as pd
import fitz
from PyPDF2 import PdfReader
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Alignment, Border, Side
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import re
from zipfile import ZipFile

def converter_pdf_para_excel(uploaded_files):
    excel_paths = []
    for nome, file in uploaded_files.items():
        reader = PdfReader(file)
        linhas_total = []
        for page in reader.pages:
            texto = page.extract_text()
            if texto:
                linhas_total += texto.splitlines()
        df = pd.DataFrame(linhas_total, columns=["Linha"])
        nome_saida = nome.replace(".pdf", "_convertido.xlsx")
        df.to_excel(nome_saida, index=False)
        excel_paths.append(nome_saida)
    return excel_paths

def preencher_marcas(lista_df, planilha_modelo):
    planilha = pd.read_excel(planilha_modelo, skiprows=2)
    planilha.columns = [col.strip() for col in planilha.columns]
    for df in lista_df:
        dados = pd.read_excel(df)
        linhas = dados["Linha"].astype(str).tolist()
        for i in range(len(linhas)-2):
            if "Fornecedor" in linhas[i] and "habilitado" in linhas[i+1] and "Marca/Fabricante:" in linhas[i+2]:
                for j in range(i, 0, -1):
                    if "Item" in linhas[j]:
                        item = int(re.findall(r"Item\s+(\d+)", linhas[j])[0])
                        marca = linhas[i+2].split(":", 1)[-1].strip()
                        planilha.loc[planilha["ITEM"] == item, "MARCA"] = marca
                        break
    return planilha

def preencher_dados_relatorio(planilha, arquivo_pdf):
   padrao_geral = re.compile(
    r"Item (\d+)[^\n]*?\n.*?"
    r"Aceito e Habilitado.*?para\s+(.*?),\s+CNPJ\s+(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}),.*?"
    r"melhor lance:\s*R\$ ([\d\.,]+).*?/ R\$ ([\d\.,]+)",
    re.DOTALL
)
        r"Aceito e Habilitado.*?para\s+(.*?),\s+CNPJ\s+(\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}),.*?"
        r"melhor lance:\s*R\$ ([\d\.,]+).*?/ R\$ ([\d\.,]+)", re.DOTALL)
    dados = {}
    with fitz.open(arquivo_pdf) as doc:
        for page in doc:
            texto = page.get_text()
            for match in padrao.findall(texto):
                item = int(match[0])
                fornecedor = match[1].strip()
                valor_unit = float(match[3].replace(".", "").replace(",", "."))
                valor_total = float(match[4].replace(".", "").replace(",", "."))
                dados[item] = {
                    "fornecedor": fornecedor,
                    "valor_unitario": valor_unit,
                    "valor_total": valor_total
                }
    for idx, row in planilha.iterrows():
        item = int(str(row["ITEM"]).replace(".0", ""))
        if item in dados:
            planilha.at[idx, "VALOR UNITÁRIO"] = dados[item]["valor_unitario"]
            planilha.at[idx, "VALOR TOTAL"] = dados[item]["valor_total"]
            planilha.at[idx, "FORNECEDOR"] = dados[item]["fornecedor"]
        else:
            planilha.at[idx, "FORNECEDOR"] = "Fracassado e/ou Deserto"
    return planilha

def gerar_atas_por_empresa(df):
    output_dir = "outputs"
    os.makedirs(output_dir, exist_ok=True)
    borda_fina = Border(left=Side(style='thin'), right=Side(style='thin'),
                        top=Side(style='thin'), bottom=Side(style='thin'))
    for fornecedor, grupo in df.groupby("FORNECEDOR"):
        nome_limpo = fornecedor[:40].replace(" ", "_").replace("/", "-")
        grupo["ITEM"] = grupo["ITEM"].astype("Int64")
        grupo_formatado = grupo[[
            "ITEM", "DESCRIÇÃO DO MATERIAL", "MARCA", "UNIDADE",
            "QUANTIDADE", "VALOR UNITÁRIO", "VALOR TOTAL"
        ]]
        wb = Workbook()
        ws = wb.active
        ws.title = "Itens Vencedores"
        ws.append(grupo_formatado.columns.tolist())
        for _, row in grupo_formatado.iterrows():
            ws.append(row.tolist())
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
            for cell in row:
                if cell.row == 1:
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                elif cell.column_letter == "B":
                    cell.alignment = Alignment(horizontal="justify", vertical="center", wrap_text=True)
                else:
                    cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = borda_fina
        for row in ws.iter_rows(min_row=2, max_row=ws.max_row):
            row[5].number_format = 'R$ #,##0.00'
            row[6].number_format = 'R$ #,##0.00'
        ws.column_dimensions["A"].width = 6
        ws.column_dimensions["B"].width = 65
        ws.column_dimensions["C"].width = 20
        ws.column_dimensions["D"].width = 20
        ws.column_dimensions["E"].width = 12
        ws.column_dimensions["F"].width = 16
        ws.column_dimensions["G"].width = 16
        excel_path = f"{output_dir}/{nome_limpo}.xlsx"
        wb.save(excel_path)

        doc = Document()
        doc.add_heading(f'Fornecedor: {fornecedor}', 0)
        table = doc.add_table(rows=1, cols=len(grupo_formatado.columns))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(grupo_formatado.columns):
            hdr_cells[i].text = str(col)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _, row in grupo_formatado.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                texto = str(value).replace(" ", "
") if grupo_formatado.columns[i] == "UNIDADE" else str(value)
                row_cells[i].text = texto
                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        docx_path = f"{output_dir}/{nome_limpo}.docx"
        doc.save(docx_path)
    return output_dir

st.title("📦 Super App - Registro de Preços")
st.markdown("Quatro etapas integradas: PDF ➜ Marcas ➜ Dados ➜ Atas")

with st.expander("1️⃣ Converter PDFs para Excel"):
    arquivos_pdf = st.file_uploader("Envie os arquivos PDF (1 por item)", type=["pdf"], accept_multiple_files=True)
    if st.button("Converter PDFs"):
        if arquivos_pdf:
            with st.spinner("Convertendo..."):
                excel_gerados = converter_pdf_para_excel({f.name: f for f in arquivos_pdf})
                st.success(f"{len(excel_gerados)} arquivos convertidos com sucesso.")
                for path in excel_gerados:
                    st.download_button("⬇ Baixar " + os.path.basename(path), open(path, "rb"), file_name=os.path.basename(path))
        else:
            st.warning("Envie ao menos um PDF.")

with st.expander("2️⃣ Preencher marcas"):
    planilha_modelo = st.file_uploader("📥 Envie a planilha modelo preenchida com os itens", type=["xlsx"], key="marcas_xlsx")
    arquivos_xlsx_convertidos = st.file_uploader("📥 Envie os arquivos .xlsx convertidos dos PDFs", type=["xlsx"], accept_multiple_files=True, key="convertidos")
    if st.button("Preencher marcas"):
        if planilha_modelo and arquivos_xlsx_convertidos:
            with st.spinner("Processando..."):
                paths = []
                for f in arquivos_xlsx_convertidos:
                    temp_path = os.path.join(tempfile.gettempdir(), f.name)
                    with open(temp_path, "wb") as out:
                        out.write(f.read())
                    paths.append(temp_path)
                planilha_saida = preencher_marcas(paths, planilha_modelo)
                path_final = "planilha_com_marcas.xlsx"
                planilha_saida.to_excel(path_final, index=False)
                st.success("Marcas preenchidas com sucesso.")
                st.download_button("⬇ Baixar planilha com marcas", open(path_final, "rb"), file_name=path_final)
        else:
            st.warning("Envie os arquivos corretamente.")

with st.expander("3️⃣ Preencher valores e fornecedores"):
    pdf_relatorio = st.file_uploader("📥 Envie o arquivo PDF do relatório de julgamento", type=["pdf"], key="pdf_rel")
    planilha_com_marcas = st.file_uploader("📥 Envie a planilha com marcas preenchidas", type=["xlsx"], key="plan_marca")
    if st.button("Preencher dados do relatório"):
        if pdf_relatorio and planilha_com_marcas:
            with st.spinner("Lendo e preenchendo..."):
                path_marca = os.path.join(tempfile.gettempdir(), planilha_com_marcas.name)
                with open(path_marca, "wb") as out:
                    out.write(planilha_com_marcas.read())
                path_pdf = os.path.join(tempfile.gettempdir(), pdf_relatorio.name)
                with open(path_pdf, "wb") as out:
                    out.write(pdf_relatorio.read())
                tabela = pd.read_excel(path_marca)
                preenchida = preencher_dados_relatorio(tabela, path_pdf)
                path_saida = "planilha_preenchida_final.xlsx"
                preenchida.to_excel(path_saida, index=False)
                st.success("Dados preenchidos.")
                st.download_button("⬇ Baixar planilha final", open(path_saida, "rb"), file_name=path_saida)
        else:
            st.warning("Envie os dois arquivos necessários.")

with st.expander("4️⃣ Gerar atas por fornecedor"):
    planilha_final = st.file_uploader("📥 Envie a planilha preenchida final", type=["xlsx"], key="atas")
    if st.button("Gerar atas"):
        if planilha_final:
            with st.spinner("Gerando atas..."):
                caminho = os.path.join(tempfile.gettempdir(), planilha_final.name)
                with open(caminho, "wb") as f:
                    f.write(planilha_final.read())
                df = pd.read_excel(caminho)
                pasta = gerar_atas_por_empresa(df)
                zip_saida = "atas_zipadas_final.zip"
                with ZipFile(zip_saida, "w") as zipf:
                    for f in os.listdir(pasta):
                        zipf.write(os.path.join(pasta, f), arcname=f)
                st.success("Atas geradas.")
                st.download_button("⬇ Baixar .zip com as atas", open(zip_saida, "rb"), file_name=zip_saida)
        else:
            st.warning("Envie a planilha final.")
